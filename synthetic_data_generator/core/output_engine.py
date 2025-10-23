"""
Multi-Format Output Engine

Outputs generated data in multiple formats:
- CSV
- JSON
- PDF
- Word (.docx)
- Excel (.xlsx)
- Markdown
"""

from typing import List, Dict, Any, Optional
from pathlib import Path
import csv
import json
from datetime import datetime


class OutputEngine:
    """Engine for outputting data in multiple formats"""

    def __init__(self):
        self.supported_formats = ['csv', 'json', 'pdf', 'word', 'excel', 'markdown']

    def export(
        self,
        data: List[Dict[str, Any]],
        output_path: str,
        format: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> str:
        """
        Export data to file in specified format

        Args:
            data: List of records to export
            output_path: Output file path (extension will be added if missing)
            format: Output format (csv, json, pdf, word, excel, markdown)
            metadata: Optional metadata to include

        Returns:
            Path to created file
        """
        format = format.lower()
        if format not in self.supported_formats:
            raise ValueError(f"Unsupported format: {format}. Supported: {self.supported_formats}")

        # Ensure correct extension
        output_path = str(output_path)
        if not output_path.endswith(f'.{format}'):
            # Handle special cases
            if format == 'word':
                output_path += '.docx'
            elif format == 'excel':
                output_path += '.xlsx'
            elif format == 'markdown':
                output_path += '.md'
            else:
                output_path += f'.{format}'

        # Export based on format
        if format == 'csv':
            return self._export_csv(data, output_path, metadata)
        elif format == 'json':
            return self._export_json(data, output_path, metadata)
        elif format == 'pdf':
            return self._export_pdf(data, output_path, metadata)
        elif format == 'word':
            return self._export_word(data, output_path, metadata)
        elif format == 'excel':
            return self._export_excel(data, output_path, metadata)
        elif format == 'markdown':
            return self._export_markdown(data, output_path, metadata)

    def _export_csv(self, data: List[Dict], output_path: str, metadata: Optional[Dict]) -> str:
        """Export to CSV format"""
        if not data:
            raise ValueError("No data to export")

        # Clean data (remove metadata fields)
        cleaned_data = [{k: v for k, v in record.items() if not k.startswith('_')}
                       for record in data]

        with open(output_path, 'w', newline='', encoding='utf-8') as f:
            if cleaned_data:
                writer = csv.DictWriter(f, fieldnames=cleaned_data[0].keys())
                writer.writeheader()
                writer.writerows(cleaned_data)

        return output_path

    def _export_json(self, data: List[Dict], output_path: str, metadata: Optional[Dict]) -> str:
        """Export to JSON format"""
        output = {
            'metadata': metadata or {
                'generated_date': datetime.now().isoformat(),
                'record_count': len(data)
            },
            'data': data
        }

        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(output, f, indent=2, ensure_ascii=False)

        return output_path

    def _export_pdf(self, data: List[Dict], output_path: str, metadata: Optional[Dict]) -> str:
        """Export to PDF format"""
        try:
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.lib import colors
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
            from reportlab.lib.units import inch
        except ImportError:
            raise ImportError("ReportLab not installed. Run: pip install reportlab")

        # Create PDF
        doc = SimpleDocTemplate(output_path, pagesize=A4)
        elements = []

        # Styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#2C3E50'),
            spaceAfter=30,
        )

        # Title
        title = metadata.get('title', 'Generated Synthetic Data') if metadata else 'Generated Synthetic Data'
        elements.append(Paragraph(title, title_style))
        elements.append(Spacer(1, 0.2*inch))

        # Metadata
        if metadata:
            meta_text = f"Generated: {metadata.get('generated_date', datetime.now().isoformat())}<br/>"
            meta_text += f"Records: {len(data)}<br/>"
            if 'geography' in metadata:
                meta_text += f"Geography: {metadata['geography']}<br/>"
            elements.append(Paragraph(meta_text, styles['Normal']))
            elements.append(Spacer(1, 0.3*inch))

        # Clean data
        cleaned_data = [{k: v for k, v in record.items() if not k.startswith('_')}
                       for record in data]

        # Table data
        if cleaned_data:
            # Headers
            headers = list(cleaned_data[0].keys())
            table_data = [headers]

            # Rows (limit to first 50 for PDF)
            for record in cleaned_data[:50]:
                row = [str(record.get(h, ''))[:30] for h in headers]  # Truncate long values
                table_data.append(row)

            # Create table
            table = Table(table_data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#3498DB')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('FONTSIZE', (0, 1), (-1, -1), 8),
            ]))

            elements.append(table)

            if len(cleaned_data) > 50:
                elements.append(Spacer(1, 0.2*inch))
                elements.append(Paragraph(f"<i>Showing first 50 of {len(cleaned_data)} records</i>",
                                        styles['Normal']))

        # Build PDF
        doc.build(elements)

        return output_path

    def _export_word(self, data: List[Dict], output_path: str, metadata: Optional[Dict]) -> str:
        """Export to Word (.docx) format"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError("python-docx not installed. Run: pip install python-docx")

        # Create document
        doc = Document()

        # Title
        title = metadata.get('title', 'Generated Synthetic Data') if metadata else 'Generated Synthetic Data'
        heading = doc.add_heading(title, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        if metadata:
            para = doc.add_paragraph()
            para.add_run(f"Generated: {metadata.get('generated_date', datetime.now().isoformat())}\n").bold = True
            para.add_run(f"Records: {len(data)}\n")
            if 'geography' in metadata:
                para.add_run(f"Geography: {metadata['geography']}\n")

        doc.add_paragraph()  # Spacer

        # Clean data
        cleaned_data = [{k: v for k, v in record.items() if not k.startswith('_')}
                       for record in data]

        # Table
        if cleaned_data:
            # Add table
            headers = list(cleaned_data[0].keys())
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Light Grid Accent 1'

            # Headers
            header_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                header_cells[i].text = str(header)
                # Bold header
                for paragraph in header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

            # Data rows (limit to 100 for Word)
            for record in cleaned_data[:100]:
                row_cells = table.add_row().cells
                for i, header in enumerate(headers):
                    row_cells[i].text = str(record.get(header, ''))

            if len(cleaned_data) > 100:
                doc.add_paragraph()
                doc.add_paragraph(f"Showing first 100 of {len(cleaned_data)} records").italic = True

        # Save
        doc.save(output_path)

        return output_path

    def _export_excel(self, data: List[Dict], output_path: str, metadata: Optional[Dict]) -> str:
        """Export to Excel (.xlsx) format"""
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill, Alignment
            from openpyxl.utils import get_column_letter
        except ImportError:
            raise ImportError("openpyxl not installed. Run: pip install openpyxl")

        # Create workbook
        wb = Workbook()

        # Data sheet
        ws_data = wb.active
        ws_data.title = "Data"

        # Clean data
        cleaned_data = [{k: v for k, v in record.items() if not k.startswith('_')}
                       for record in data]

        if cleaned_data:
            # Headers
            headers = list(cleaned_data[0].keys())
            for col, header in enumerate(headers, 1):
                cell = ws_data.cell(row=1, column=col, value=header)
                cell.font = Font(bold=True, color="FFFFFF")
                cell.fill = PatternFill(start_color="3498DB", end_color="3498DB", fill_type="solid")
                cell.alignment = Alignment(horizontal="center")

            # Data rows
            for row_idx, record in enumerate(cleaned_data, 2):
                for col_idx, header in enumerate(headers, 1):
                    value = record.get(header, '')
                    ws_data.cell(row=row_idx, column=col_idx, value=value)

            # Adjust column widths
            for col in range(1, len(headers) + 1):
                ws_data.column_dimensions[get_column_letter(col)].width = 15

        # Metadata sheet
        if metadata:
            ws_meta = wb.create_sheet("Metadata")
            ws_meta['A1'] = "Property"
            ws_meta['B1'] = "Value"
            ws_meta['A1'].font = Font(bold=True)
            ws_meta['B1'].font = Font(bold=True)

            row = 2
            for key, value in metadata.items():
                ws_meta[f'A{row}'] = str(key)
                ws_meta[f'B{row}'] = str(value)
                row += 1

        # Save
        wb.save(output_path)

        return output_path

    def _export_markdown(self, data: List[Dict], output_path: str, metadata: Optional[Dict]) -> str:
        """Export to Markdown format"""
        lines = []

        # Title
        title = metadata.get('title', 'Generated Synthetic Data') if metadata else 'Generated Synthetic Data'
        lines.append(f"# {title}\n")

        # Metadata
        if metadata:
            lines.append(f"**Generated:** {metadata.get('generated_date', datetime.now().isoformat())}  ")
            lines.append(f"**Records:** {len(data)}  ")
            if 'geography' in metadata:
                lines.append(f"**Geography:** {metadata['geography']}  ")
            lines.append("")

        # Clean data
        cleaned_data = [{k: v for k, v in record.items() if not k.startswith('_')}
                       for record in data]

        # Table
        if cleaned_data:
            lines.append("## Data\n")

            # Headers
            headers = list(cleaned_data[0].keys())
            lines.append("| " + " | ".join(headers) + " |")
            lines.append("|" + "|".join(["---" for _ in headers]) + "|")

            # Rows (limit to 50 for readability)
            for record in cleaned_data[:50]:
                row = [str(record.get(h, ''))[:50] for h in headers]  # Truncate long values
                lines.append("| " + " | ".join(row) + " |")

            if len(cleaned_data) > 50:
                lines.append(f"\n*Showing first 50 of {len(cleaned_data)} records*")

        # Write file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(lines))

        return output_path


class BatchOutputEngine:
    """Export to multiple formats at once"""

    def __init__(self):
        self.engine = OutputEngine()

    def export_multiple(
        self,
        data: List[Dict[str, Any]],
        output_dir: str,
        base_name: str,
        formats: List[str],
        metadata: Optional[Dict[str, Any]] = None
    ) -> Dict[str, str]:
        """
        Export data to multiple formats

        Args:
            data: Data to export
            output_dir: Output directory
            base_name: Base filename (without extension)
            formats: List of formats to export
            metadata: Optional metadata

        Returns:
            Dictionary mapping format to output path
        """
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)

        results = {}

        for fmt in formats:
            output_path = output_dir / f"{base_name}.{fmt}"
            try:
                path = self.engine.export(data, str(output_path), fmt, metadata)
                results[fmt] = path
                print(f"✓ Exported to {path}")
            except Exception as e:
                print(f"✗ Failed to export to {fmt}: {e}")
                results[fmt] = None

        return results
